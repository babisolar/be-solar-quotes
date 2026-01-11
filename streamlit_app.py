import streamlit as st
from docx import Document
from datetime import datetime
from num2words import num2words
import os
import io
import json

# ---------------- CONFIG ----------------

APP_TITLE = "Babi Enterprise Solar Installation Quotation"
USERNAME = "besolar"
PASSWORD = "solar@2025"
RATE_PER_KW = 70000

DATA_DIR = "data"
os.makedirs(DATA_DIR, exist_ok=True)

INVOICE_COUNTER_FILE = os.path.join(DATA_DIR, "invoice_counter.json")
AGREEMENT_COUNTER_FILE = os.path.join(DATA_DIR, "agreement_counter.json")

FOOTER_TEXT = """
Babi Enterprise  
Khowang, Dibrugarh, Assam – 785676  
Phone: 9678244548 / 9678689212  
Email: udaskhowang@gmail.com / neelbaruah@gmail.com
"""

WHATSAPP_TEMPLATE = """Dear {name},

Your solar quotation and agreement are ready. Please find the attached documents.

Thank you for choosing Babi Enterprise.
"""

# ---------------- LOGIN ----------------

def login():
    st.title(APP_TITLE)
    st.subheader("Login")

    username = st.text_input("Username")
    password = st.text_input("Password", type="password")

    if st.button("Login"):
        if username == USERNAME and password == PASSWORD:
            st.session_state.logged_in = True
            st.rerun()
        else:
            st.error("Invalid credentials")

# ---------------- UTILITIES ----------------

def format_legal_date(dt):
    day = dt.day
    if 11 <= day <= 13:
        suffix = "th"
    else:
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    return f"{day}{suffix} Day of {dt.strftime('%B %Y')}"

def replace_everywhere(doc, data):
    for para in doc.paragraphs:
        for k, v in data.items():
            if k in para.text:
                para.text = para.text.replace(k, v)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for k, v in data.items():
                    if k in cell.text:
                        cell.text = cell.text.replace(k, v)

# ---------------- INVOICE NUMBER LOGIC ----------------

def get_next_invoice_ref():
    now = datetime.now()
    key = now.strftime("%m/%y")

    if os.path.exists(INVOICE_COUNTER_FILE):
        with open(INVOICE_COUNTER_FILE, "r") as f:
            data = json.load(f)
    else:
        data = {}

    if key not in data:
        data[key] = 1
    else:
        data[key] += 1

    with open(INVOICE_COUNTER_FILE, "w") as f:
        json.dump(data, f)

    return f"BE/KNG/PMSG/QTN/{key}/{str(data[key]).zfill(4)}"

# ---------------- AGREEMENT NUMBER LOGIC ----------------

def get_next_agreement_no():
    year = str(datetime.now().year)

    if os.path.exists(AGREEMENT_COUNTER_FILE):
        with open(AGREEMENT_COUNTER_FILE, "r") as f:
            data = json.load(f)
    else:
        data = {}

    if year not in data:
        data[year] = 1
    else:
        data[year] += 1

    with open(AGREEMENT_COUNTER_FILE, "w") as f:
        json.dump(data, f)

    return f"AG/SG/APDCL/{year}/{str(data[year]).zfill(4)}"

# ---------------- MAIN APP ----------------

def main_app():
    st.title(APP_TITLE)

    st.subheader("Customer Details")

    name = st.text_input("Customer Name")
    phone = st.text_input("Phone Number")
    address = st.text_area("Address")
    consumer_no = st.text_input("APDCL Consumer Number")
    subdivision = st.text_input("Subdivision")

    capacity = st.selectbox("System Capacity (kW)", [3, 4.5, 5, 10])

    phase = "Three Phase" if capacity >= 5 else "Single Phase"
    total_amount = int(capacity * RATE_PER_KW)
    amount_words = f"Rupees {num2words(total_amount, lang='en_IN').replace('-', ' ')} only"

    st.info(f"Phase: {phase}")
    st.success(f"Total Amount: ₹{total_amount:,}")

    w1_name = st.text_input("Witness 1 Name")
    w1_phone = st.text_input("Witness 1 Phone")
    w2_name = st.text_input("Witness 2 Name")
    w2_phone = st.text_input("Witness 2 Phone")

    if st.button("Generate Documents"):
        if not name or not phone:
            st.error("Customer name and phone are required.")
            return

        now = datetime.now()
        invoice_date = now.strftime("%d/%m/%Y")
        agreement_date = format_legal_date(now)
        agreement_no = get_next_agreement_no()
        ref_no = get_next_invoice_ref()

        data = {
            "{{REF_NO}}": ref_no,
            "{{DATE}}": invoice_date,
            "{{AGREEMENT_NO}}": agreement_no,
            "{{AGREEMENT_DATE}}": agreement_date,
            "{{CUSTOMER_NAME}}": name,
            "{{ADDRESS}}": address,
            "{{PHONE}}": phone,
            "{{CONSUMER_NO}}": consumer_no,
            "{{SYSTEM_CAPACITY}}": f"{capacity} kW",
            "{{TOTAL_AMOUNT}}": f"{total_amount:,}",
            "{{SYSTEM_COST}}": f"{total_amount:,}",
            "{{AMOUNT_IN_WORDS}}": amount_words,
            "{{PHASE_TYPE}}": phase,
            "{{APDCL_SUBDIVISION}}": subdivision,
            "{{W1_NAME}}": w1_name,
            "{{W1_PHONE}}": w1_phone,
            "{{W2_NAME}}": w2_name,
            "{{W2_PHONE}}": w2_phone,
        }

        invoice_doc = Document("Invoice Sample.docx")
        replace_everywhere(invoice_doc, data)
        invoice_buffer = io.BytesIO()
        invoice_doc.save(invoice_buffer)
        invoice_buffer.seek(0)

        agreement_doc = Document("Agreement.docx")
        replace_everywhere(agreement_doc, data)
        agreement_buffer = io.BytesIO()
        agreement_doc.save(agreement_buffer)
        agreement_buffer.seek(0)

        st.success("Documents generated successfully")

        st.download_button(
            label="Download Invoice DOCX",
            data=invoice_buffer,
            file_name=f"Invoice_{name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        st.download_button(
            label="Download Agreement DOCX",
            data=agreement_buffer,
            file_name=f"{agreement_no.replace('/', '_')}_{name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        whatsapp_msg = WHATSAPP_TEMPLATE.format(name=name)
        encoded_msg = whatsapp_msg.replace("\n", "%0A").replace(" ", "%20")
        whatsapp_url = f"https://wa.me/?text={encoded_msg}"

        st.markdown(f"📲 [Share on WhatsApp]({whatsapp_url})")

    st.markdown("---")
    st.markdown(FOOTER_TEXT)

# ---------------- ROUTER ----------------

if "logged_in" not in st.session_state:
    st.session_state.logged_in = False

if not st.session_state.logged_in:
    login()
else:
    main_app()
